import traceback
try:
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    # Helper to set font
    def set_font(run, name, size, bold=False):
        run.font.name = name
        run.font.size = Pt(size)
        run.font.bold = bold
        r = run._element
        r.rPr.rFonts.set(qn('w:eastAsia'), name)

    doc = Document()

    # Title Page
    p = doc.add_paragraph()
    run = p.add_run("QGIS-FME Form Connector Algorithm")
    set_font(run, 'Arial', 16, bold=True)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run("Author: GIS Innovation Sdn. Bhd. (Lyes)")
    set_font(run, 'Arial', 11)
    p = doc.add_paragraph()
    run = p.add_run("Date: 2025-04-23")
    set_font(run, 'Arial', 11)
    p = doc.add_paragraph()
    run = p.add_run("Version: 0.3")
    set_font(run, 'Arial', 11)
    p = doc.add_paragraph()
    run = p.add_run("Copyright: (C) 2025 by GIS Innovation Sdn. Bhd.")
    set_font(run, 'Arial', 11)

    # Page break
    from docx.enum.section import WD_SECTION
    section = doc.sections[-1]
    section.start_type = WD_SECTION.NEW_PAGE

    doc.add_paragraph('Table of Contents', style='Heading 1')
    doc.add_paragraph('1. Introduction')
    doc.add_paragraph('2. Overview of the Algorithm')
    doc.add_paragraph('3. Key Classes and Components')
    doc.add_paragraph('4. User Interface and Workflow')
    doc.add_paragraph('5. Usage Instructions')
    doc.add_paragraph('6. Error Handling and Logging')
    doc.add_paragraph('7. Requirements & Dependencies')
    doc.add_paragraph('8. Example Usage')
    doc.add_paragraph('9. Contact & Support')

    # Introduction
    p = doc.add_paragraph()
    run = p.add_run("1. Introduction")
    set_font(run, 'Arial', 14, bold=True)
    run = doc.add_paragraph().add_run(
        "The QGIS-FME Form Connector Algorithm is a QGIS Processing Algorithm plugin that enables seamless integration "
        "between QGIS and FME workspaces. Users can browse, configure, and execute FME transformations directly from QGIS, "
        "making advanced spatial data workflows more accessible.")
    set_font(run, 'Arial', 11)

    # Overview
    p = doc.add_paragraph()
    run = p.add_run("2. Overview of the Algorithm")
    set_font(run, 'Arial', 14, bold=True)
    run = doc.add_paragraph().add_run(
        "Purpose:\nTo allow users to select FME workspaces (.fmw), configure parameters, and run FME transformations from within QGIS.\n"
        "Features:\n- Browse and select FME workspaces.\n- View and edit workspace parameters.\n- Configure and run FME command-line executions.\n- Load transformation results back into QGIS.\n- Comprehensive error handling and logging.")
    set_font(run, 'Arial', 11)

    # Key Classes
    p = doc.add_paragraph()
    run = p.add_run("3. Key Classes and Components")
    set_font(run, 'Arial', 14, bold=True)
    run = doc.add_paragraph().add_run(
        "QGISFMEFormAlgorithmAlgorithm (Main Algorithm Class):\n"
        "- Inherits from QgsProcessingAlgorithm.\n"
        "- Handles the main logic for parameter definition, command construction, and execution.\n"
        "- Key methods: initAlgorithm(), processAlgorithm(), icon(), svgIconPath(), name(), displayName().\n\n"
        "CustomParametersWidget:\n"
        "- Inherits from WidgetWrapper.\n"
        "- Provides a custom UI for parameter entry and command editing.\n\n"
        "FMEFileLister:\n"
        "- Custom QWidget for browsing and selecting FME workspaces.\n\n"
        "CollapsibleGroupBox:\n"
        "- Custom QGroupBox with collapsible/expandable sections for better UI organization.")
    set_font(run, 'Arial', 11)

    # UI and Workflow
    p = doc.add_paragraph()
    run = p.add_run("4. User Interface and Workflow")
    set_font(run, 'Arial', 14, bold=True)
    run = doc.add_paragraph().add_run(
        "- Workspace Selection: Use the file lister to browse directories and select an FME workspace (.fmw).\n"
        "- Parameter Configuration: Parameters from the selected workspace are displayed in a table. Users can edit values as needed.\n"
        "- Command Editing: The FME command line is shown in a plain text box for advanced users to review or modify.\n"
        "- Execution: Click 'Run' to execute the FME transformation. Outputs are saved to temporary files and can be loaded into QGIS.")
    set_font(run, 'Arial', 11)

    # Usage Instructions
    p = doc.add_paragraph()
    run = p.add_run("5. Usage Instructions")
    set_font(run, 'Arial', 14, bold=True)
    run = doc.add_paragraph().add_run(
        "1. Install the Plugin: Place the plugin in your QGIS plugins directory and enable it.\n"
        "2. Open the Algorithm: Find 'FME Form Connector Algorithm' in the QGIS Processing Toolbox.\n"
        "3. Select Workspace: Use the UI to browse and select your .fmw file.\n"
        "4. Configure Parameters: Edit any required parameters in the table.\n"
        "5. Run the Algorithm: Review or modify the command, then execute.\n"
        "6. View Results: Output layers can be loaded directly into QGIS.")
    set_font(run, 'Arial', 11)

    # Error Handling
    p = doc.add_paragraph()
    run = p.add_run("6. Error Handling and Logging")
    set_font(run, 'Arial', 14, bold=True)
    run = doc.add_paragraph().add_run(
        "- Detailed error handling using Python’s traceback module.\n"
        "- Errors are reported with line numbers and messages in QGIS message logs.\n"
        "- User-friendly error dialogs (QMessageBox) for critical issues.")
    set_font(run, 'Arial', 11)

    # Requirements
    p = doc.add_paragraph()
    run = p.add_run("7. Requirements & Dependencies")
    set_font(run, 'Arial', 14, bold=True)
    run = doc.add_paragraph().add_run(
        "- QGIS 3.x\n- FME Desktop (with command-line access)\n- Python packages: PyQt5, qgis.core, win32api, etc.")
    set_font(run, 'Arial', 11)

    # Example Usage
    p = doc.add_paragraph()
    run = p.add_run("8. Example Usage")
    set_font(run, 'Arial', 14, bold=True)
    run = doc.add_paragraph().add_run(
        "1. Select an FME workspace.\n"
        "2. The plugin reads parameters and allows editing.\n"
        "3. The command is constructed (e.g., \"C:\\Program Files\\FME\\fme.exe\" \"workspace.fmw\" --SourceDataset_GEOJSON \"input.geojson\" --DestDataset_GEOJSON \"output.geojson\").\n"
        "4. Run the command and view the output in QGIS.")
    set_font(run, 'Arial', 11)

    # Contact
    p = doc.add_paragraph()
    run = p.add_run("9. Contact & Support")
    set_font(run, 'Arial', 14, bold=True)
    run = doc.add_paragraph().add_run(
        "For support or further information, contact:\n"
        "GIS Innovation Sdn. Bhd.\n"
        "Email: your-support-email@example.com\n"
        "Author: Lyes")
    set_font(run, 'Arial', 11)

    # Save the document
    doc.save('QGIS_FME_Form_Connector_Algorithm_Documentation.docx')
except Exception as e:
    with open('docx_generation_error.txt', 'w', encoding='utf-8') as f:
        f.write('An error occurred while generating the documentation:\n')
        f.write(str(e) + '\n')
        import traceback
        f.write(traceback.format_exc())